#!/usr/bin/env python3
"""
Builds the updated Nabholz Appraisal Write-Up Consistency Check SOP as a .docx file.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ─── Styles ────────────────────────────────────────────────────────────────

def set_heading(para, level=1):
    para.style = f'Heading {level}'

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p

def add_para(doc, text, bold=False, italic=False, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_table_with_header(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
        cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri+1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            if isinstance(val, tuple):
                # (text, bold)
                run = cell.paragraphs[0].add_run(val[0])
                run.bold = val[1]
            else:
                cell.text = str(val)
    return table

# ─── Document ──────────────────────────────────────────────────────────────

# Title
title = doc.add_heading('Nabholz Appraisal Write-Up Consistency Check SOP', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Knowledge File for Custom GPT: Appraisal Write-Up Consistency Checker\nVersion 1.1 — Front-End Write-Up Review (Updated May 2026)')
r.italic = True

doc.add_paragraph()

# ─── Plain-English Purpose ─────────────────────────────────────────────────
add_heading(doc, 'Plain-English Purpose', 1)
add_para(doc, 'This SOP tells the Custom GPT how to compare the initial appraisal report write-up against the order packet and source documents. The goal is to catch front-end setup problems before the file moves further into production.')

# ─── Section 1 ─────────────────────────────────────────────────────────────
add_heading(doc, '1. Purpose', 1)
add_para(doc, 'The Appraisal Write-Up Consistency Checker identifies factual inconsistencies, missing information, and items requiring staff review after the appraisal report has been initially written up.')
add_para(doc, 'It is designed to catch front-end errors such as:')
for item in [
    'Wrong county, address, or ZIP',
    'Acreage conflicts and parcel issues',
    'Borrower/buyer name conflicts',
    'Contract price mismatches',
    'FHA/VA case number issues',
    'Unsigned or missing contracts',
    'HOA/PUD conflicts',
    'Seller name vs. tax card owner mismatches',
    'Non-arms-length transaction indicators',
    'GP orders missing intended use or intended user',
    'Report fee outside expected schedule',
    'Due date calculation errors',
    'Possible investment property / rent schedule requirements',
]:
    add_bullet(doc, item)
add_para(doc, 'This GPT does not replace staff review, appraiser judgment, or final quality control.')

# ─── Section 2 ─────────────────────────────────────────────────────────────
add_heading(doc, '2. Scope of Version 1.1', 1)
add_para(doc, 'Version 1.1 is limited to the front-end write-up consistency check, run after the report has been initially written up and before the file moves deeper into production.')

headers = ['Included in Version 1.1', 'Excluded From Version 1.1']
rows = [
    ('Property identity consistency', 'Comparable selection'),
    ('Assignment and order setup consistency', 'Value opinions'),
    ('Contract and transaction data consistency', 'Adjustment support or market analysis'),
    ('Site, acreage, parcel, and legal-description conflicts', 'GLA verification before inspection'),
    ('Year built and HOA/PUD consistency', 'Zoning review'),
    ('Basic FHA/VA case number checks', 'Full FHA/VA/RD/Fannie/lender compliance review'),
    ('FHA new construction builder certification presence', 'Final signature/inspector/appraiser logic'),
    ('Contract execution (fully signed by both parties)', ''),
    ('Non-arms-length transaction detection', ''),
    ('Seller name vs. tax card owner check', ''),
    ('GP order intended use/user verification', ''),
    ('Report fee reasonableness check', ''),
    ('Due date reasonableness check', ''),
    ('Investment property / rent schedule flag', ''),
]
add_table_with_header(doc, headers, rows)

# ─── Section 3 ─────────────────────────────────────────────────────────────
add_heading(doc, '3. Timing and Responsibility', 1)
headers = ['Item', 'Rule']
rows = [
    ('When this check is run', 'After Anna completes the initial report write-up.'),
    ('Who runs the check', 'Anna.'),
    ('Who reviews the result', 'Anna reviews and resolves all issues she can resolve.'),
    ('Escalation', 'If Anna cannot resolve an issue, she requests missing information from the client and/or escalates to Amy.'),
    ('Red flag handling', 'Red flags must be corrected, documented, or escalated before the file moves forward.'),
    ('Rerun rule', 'Rerun the GPT after any corrections are made.'),
]
add_table_with_header(doc, headers, rows)

# ─── Section 4 ─────────────────────────────────────────────────────────────
add_heading(doc, '4. Required Upload Packet', 1)
add_para(doc, 'Upload the available documents listed below. If a document is missing, the GPT will still run the check but will clearly state what could not be verified.')
headers = ['Document', 'When Required / Expected']
rows = [
    ('Report front page', 'Required every time.'),
    ('Order form / engagement letter', 'Required every time.'),
    ('Sales contract', 'Required for purchase assignments when available.'),
    ('MLS subject sheet', 'Expected when available.'),
    ('Tax card / assessor record', 'Required every time when available.'),
    ('Parcel map, plat, legal description, deed, survey, or other parcel/legal support', 'Required when available and relevant.'),
    ('Seller disclosure', 'Include when available — used for utility and property condition flags.'),
]
add_table_with_header(doc, headers, rows)

# ─── Section 5 ─────────────────────────────────────────────────────────────
add_heading(doc, '5. File Naming', 1)
add_para(doc, 'Use clear file names before uploading. Names do not have to be perfect, but they should make the document type obvious.')
headers = ['Document Type', 'Recommended File Name']
rows = [
    ('Report front page', '01 Report Front Page.pdf'),
    ('Order form / engagement letter', '02 Order Form.pdf'),
    ('Sales contract', '03 Sales Contract.pdf'),
    ('Contract amendments/addenda', '03B Contract Amendment.pdf, 03C Repair Addendum.pdf'),
    ('MLS subject sheet', '04 MLS Subject.pdf'),
    ('Tax card / assessor record', '05 Tax Card.pdf'),
    ('Multiple tax cards', '05A Tax Card Parcel 1.pdf, 05B Tax Card Parcel 2.pdf'),
    ('Parcel/legal support', '06 Parcel Map.pdf  or  06 Legal Description.pdf'),
    ('Seller disclosure', '07 Seller Disclosure.pdf'),
]
add_table_with_header(doc, headers, rows)

# ─── Section 6 ─────────────────────────────────────────────────────────────
add_heading(doc, '6. Field Checklist and Check Rules', 1)
add_para(doc, 'The GPT extracts and compares the following fields when they appear in the uploaded documents. If the source needed for a field is missing, the GPT marks the field as Unable to Verify instead of OK.')

# 6A
add_heading(doc, 'A. Property Identity', 2)
headers = ['Field', 'Compare Across']
rows = [
    ('Street address', 'Report front page, order, contract, MLS, tax card, parcel/legal documents'),
    ('City', 'Report front page, order, contract, MLS, tax card'),
    ('County', 'Report front page, order, MLS, tax card, parcel/legal documents'),
    ('ZIP code', 'Report front page, order, contract, MLS, tax card'),
    ('Parcel number / APN / tax ID', 'Report front page, order if shown, MLS, tax card, parcel/legal documents'),
    ('Legal description', 'Report front page, contract, tax card, parcel/legal documents'),
    ('Number of parcels', 'Report front page, order, contract, tax card, parcel/legal documents'),
    ('Subdivision / lot / block', 'Report front page, MLS, tax card, legal documents when available'),
]
add_table_with_header(doc, headers, rows)

# 6B
add_heading(doc, 'B. Assignment Setup', 2)
headers = ['Field', 'Compare Across / Rule']
rows = [
    ('Client/lender name', 'Report front page and order.'),
    ('AMC / portal name', 'Order and internal documents when shown.'),
    ('Borrower name', 'Report front page, order, and contract when applicable.'),
    ('Buyer name', 'Report front page, order, and contract for purchase assignments.'),
    ('Seller name', 'Contract, MLS, tax card when applicable.'),
    ('Owner name', 'Tax card, MLS, and contract when shown.'),
    ('Assignment type', 'Report/front-end setup and order when shown.'),
    ('FHA case number', 'Compare report front page to order/source documents when applicable.'),
    ('VA case number', 'Compare report front page to order/source documents when applicable.'),
    ('Due date', 'Order and internal setup when shown. See Due Date Reasonableness check below.'),
]
add_table_with_header(doc, headers, rows)
add_para(doc, 'Do not check assigned appraiser or inspector name. The inspector may not be assigned at the write-up stage.')
add_para(doc, 'Do not treat the report front page as the reliable source for loan type. For FHA and VA assignments, focus on whether the case number matches the order/source documents.')

# 6C
add_heading(doc, 'C. Contract / Transaction Data', 2)
headers = ['Field', 'Compare Across / Rule']
rows = [
    ('Contract price', 'Report front page, order, contract, and MLS when available.'),
    ('Contract date', 'Report front page if shown and contract.'),
    ('Buyer names', 'Report front page, order, and contract.'),
    ('Seller names', 'Contract, MLS, and tax card when applicable.'),
    ('Seller concessions', 'Contract, MLS, and report if addressed.'),
    ('Personal property', 'Contract, MLS, and report if addressed.'),
    ('Repairs or completion items', 'Contract, amendments/addenda, and report if addressed.'),
    ('Site size being conveyed', 'Contract, MLS, report, tax card, and parcel/legal documents.'),
    ('Included or excluded parcels', 'Contract, parcel/legal documents, and report.'),
]
add_table_with_header(doc, headers, rows)

p = doc.add_paragraph()
r = p.add_run('Contract Execution Check: ')
r.bold = True
p.add_run('Verify the sales contract is fully executed — signed by both buyer and seller. If not signed by both parties → Red Flag. If only a counter offer or addendum was uploaded without a fully executed main contract → Red Flag. Recommended action: obtain fully executed contract before the file moves forward.')

p = doc.add_paragraph()
r = p.add_run('Non-Arms-Length Detection: ')
r.bold = True
p.add_run('If the buyer and seller share the same last name, or either name appears on both sides of the transaction → Yellow Flag: "Buyer and seller may be related — confirm arms-length transaction or note as non-arms-length in the report."')

# 6D
add_heading(doc, 'D. Site / Parcel / Acreage', 2)
headers = ['Field', 'Compare Across / Rule']
rows = [
    ('Reported site size', 'Report front page, order, contract, MLS, tax card, and legal docs.'),
    ('Tax record acreage', 'Tax card / assessor record.'),
    ('MLS acreage', 'MLS subject sheet when available.'),
    ('Contract acreage', 'Sales contract when applicable.'),
    ('Survey/deed/legal acreage', 'Survey, deed, legal description, plat, or parcel docs when available.'),
    ('Partial conveyance language', 'Contract, order, legal docs, deed, or survey when available.'),
    ('Multiple parcels', 'Tax card, parcel map, legal docs, contract, and report.'),
    ('Parcel map consistency', 'Parcel map, tax card, report address/legal info.'),
]
add_table_with_header(doc, headers, rows)
add_para(doc, 'Acreage and Parcel Rule: Acreage, parcel count, legal description, partial conveyance, and multi-parcel issues are high priority. Do not assume tax acreage is correct or incorrect. Tax records may show a parent parcel while the contract or MLS may show only the acreage being conveyed. Flag the conflict and recommend staff review.')

# 6E
add_heading(doc, 'E. Basic Property Data', 2)
headers = ['Field', 'Rule']
rows = [
    ('Year built', 'Compare report, MLS, and tax card when available.'),
    ('HOA/PUD indication', 'If MLS shows an HOA or association fee but the PUD box is not checked on the report → Red Flag. Otherwise compare report, MLS, contract/order when shown.'),
]
add_table_with_header(doc, headers, rows)
add_para(doc, 'Do not check GLA, zoning, occupancy, or property type unless an obvious front-end setup conflict is visible from the uploaded documents.')

# 6F
add_heading(doc, 'F. Basic Utility Consistency', 2)
headers = ['Field', 'Rule']
rows = [
    ('Public water vs. private well', 'Compare report front page against MLS, seller disclosure, tax/assessor, or other uploaded sources. Clear conflict → Red Flag. Unclear → Yellow Flag or Unable to Verify.'),
    ('Public sewer vs. septic', 'Same rule as water source above.'),
    ('Other obvious utility conflicts', 'Flag if clearly shown in uploaded documents.'),
]
add_table_with_header(doc, headers, rows)
add_para(doc, 'Do not perform a full utility compliance review.')

# 6G
add_heading(doc, 'G. Basic Agency Setup (FHA / VA / RD)', 2)
headers = ['Field / Item', 'Rule']
rows = [
    ('FHA case number', 'Compare report front page to order/source documents when applicable.'),
    ('VA case number', 'Compare report front page to order/source documents when applicable.'),
    ('FHA new construction builder certification', 'Check document presence when FHA new construction is indicated.'),
]
add_table_with_header(doc, headers, rows)
add_para(doc, 'Do not perform a full FHA, VA, RD, Fannie Mae, or lender compliance review. Only flag obvious front-end setup conflicts or missing agency documents visible from the uploaded packet.')

# 6H — NEW
add_heading(doc, 'H. Seller Name vs. Tax Card Owner', 2)
add_para(doc, 'Compare the seller name on the sales contract against the property owner on the tax card.')
add_bullet(doc, 'Minor variation (typo, middle name, married name) → Yellow Flag.')
add_bullet(doc, 'Completely different name or business entity → Red Flag. Recommended action: Check Arkansas Secretary of State at ark.org/corp-search to verify the signer is a registered agent of the business. If confirmed, save as "Proof of Ownership" PDF. If no relationship found, request Proof of Ownership from client and leave a Tracker Update.')

# 6I — NEW
add_heading(doc, 'I. Rent Schedule / Investment Property', 2)
add_para(doc, 'If any uploaded document mentions "1007 form," "investment property," "rental income," or "rent schedule" → Yellow Flag: "Docs indicate possible investment property or rent schedule requirement — confirm with Amy and request current rental rates from client early in the process if applicable."')

# 6J — NEW
add_heading(doc, 'J. Due Date Reasonableness', 2)
headers = ['Order Type', 'Expected Due Date Rule']
rows = [
    ('Lender / conventional orders', 'Tracker due date should be 1 business day before the date listed on the order or engagement letter.'),
    ('VA orders', 'Due date should be approximately 7 business days after the appraisal request date.'),
    ('GP (private) orders', 'Due date should be approximately 7 business days after the appraisal request date.'),
]
add_table_with_header(doc, headers, rows)
add_para(doc, 'If the due date appears to match the lender\'s listed due date exactly (not 1 day prior) → Yellow Flag. If VA/GP due date does not appear to be approximately 7 business days from request date → Yellow Flag. If order type is unclear → Unable to Verify, skip this check.')

# 6K — NEW
add_heading(doc, 'K. Report Fee Reasonableness', 2)
add_para(doc, 'If a report fee is visible in the uploaded documents, compare against the Nabholz standard fee schedule:')
headers = ['Order Type', 'Standard Fee']
rows = [
    ('Lender / conventional', '$600'),
    ('VA — single-family residential', '$700'),
    ('VA — manufactured home', '$750'),
    ('VA — 2–4 unit multi-family', '$800'),
    ('GP (private / non-lender)', '$550'),
]
add_table_with_header(doc, headers, rows)
add_para(doc, 'Fee outside expected range without explanation → Yellow Flag: "Report fee appears outside standard schedule — verify against order details."')

# 6L — NEW
add_heading(doc, 'L. GP Orders — Intended Use and Intended User', 2)
add_para(doc, 'For GP (private / non-lender) assignments, verify that both intended use and intended user are confirmed in the order or engagement documentation.')
add_para(doc, 'If either is missing → Red Flag: "GP orders require confirmed intended use and intended user before acceptance — escalate to Amy."')

# ─── Section 7 ─────────────────────────────────────────────────────────────
add_heading(doc, '7. Severity Rules', 1)
headers = ['Status', 'Meaning']
rows = [
    ('🔴 Red Flag', 'Issue could cause the report to be materially wrong or needs to stop until resolved.'),
    ('🟡 Yellow Flag', 'Issue may be explainable but Anna needs to review it.'),
    ('ℹ️ Information Only', 'Useful note, but not necessarily a problem.'),
    ('✅ OK', 'No meaningful inconsistency found and relevant source documents were uploaded.'),
    ('⚪ Unable to Verify', 'The necessary source document or readable information was not provided.'),
]
add_table_with_header(doc, headers, rows)

add_heading(doc, 'Red Flag Examples', 2)
for item in [
    'County mismatch',
    'Address mismatch',
    'Parcel number mismatch',
    'Legal description materially different',
    'Contract price mismatch',
    'FHA or VA case number mismatch',
    'MLS/contract acreage conflicts with tax/parcel/legal documents',
    'Multiple parcels shown in source documents but not reflected in the report',
    'Partial conveyance language present but not addressed',
    'Buyer/borrower appears to be a different person',
    'Required FHA new construction builder certification missing',
    'Contract not fully executed (missing buyer or seller signature)',
    'HOA/PUD — MLS shows HOA fee but PUD box not checked on report',
    'Seller name completely unrelated to tax card owner or business entity not verified',
    'Clear public sewer/septic or public water/well conflict',
    'GP order missing intended use or intended user',
]:
    add_bullet(doc, item)

add_heading(doc, 'Yellow Flag Examples', 2)
for item in [
    'Minor buyer/borrower name variation',
    'Seller name slightly differs from tax owner (typo, married name)',
    'Year built differs between MLS/tax/report',
    'HOA/PUD information appears in one source but is unclear',
    'Tax acreage differs from contract/MLS (possible parent parcel issue)',
    'Legal description abbreviated in one source',
    'Seller concessions present and need confirmation in the report',
    'Personal property mentioned in contract/MLS',
    'Repairs or completion items mentioned in contract/addendum',
    'Non-arms-length indicator (buyer and seller share last name)',
    'Due date does not follow the expected calculation rule',
    'Report fee outside expected range without explanation',
    'Possible investment property or rent schedule requirement',
    'Counter offer or addendum present without confirmed fully executed main contract',
    'Unclear utility information',
]:
    add_bullet(doc, item)

add_heading(doc, 'Information Only Examples', 2)
for item in [
    'Tax card rounds acreage differently',
    'MLS uses shortened subdivision name',
    'Due date shown in order',
    'Parcel map included and appears consistent',
    'Contract has no concessions noted',
    'MLS mentions a shed or minor feature with no conflict',
]:
    add_bullet(doc, item)

# ─── Section 8 ─────────────────────────────────────────────────────────────
add_heading(doc, '8. Required Output Format', 1)
add_para(doc, 'The GPT uses the following output structure every time. Output should be concise, direct, and practical for Anna to use.')

output_template = """Documents Uploaded:
  ✅ Order form / engagement letter
  ✅ Tax card
  ✅ Sales contract
  ⚪ MLS subject sheet — not uploaded
  ⚪ Parcel map — not uploaded
(List all expected document types and mark each as uploaded or not uploaded.)

─────────────────────────────────────────

QC RESULT: PASS / REVIEW REQUIRED / STOP - OFFICE REVIEW REQUIRED

Property:
[Address]

Summary:
[Short plain-English summary of what was found.]

Red Flags:
1. [Issue]
   - Source conflict:
   - Recommended action:

Yellow Flags:
1. [Issue]
   - Source conflict:
   - Recommended action:

Discrepancy Grid (Red and Yellow Flags only):
| Status | Field | Report Value | Source Value(s) | Issue | Recommended Action |
  🔴 Red Flag   🟡 Yellow Flag   Bold conflicting values   Keep rows short.

ℹ️ Information Only:
  • [Note]

✅ Verified / OK Items:
  • [Important OK items only — not every possible field]

⚪ Missing Documents / Unable to Verify:
  • [Missing item — brief impact note]

Recommended Staff Actions:
1. [Action]
2. [Action]

Suggested Report Comment:
[Only include if a Red or Yellow Flag likely requires a report explanation.
 Do not write a comment for items already consistent and properly addressed.]

─────────────────────────────────────────

Overall result labels:
  QC RESULT: PASS                    — no red or yellow flags found
  QC RESULT: REVIEW REQUIRED        — yellow flags found, no red flags
  QC RESULT: STOP - OFFICE REVIEW   — any red flag found"""

p = doc.add_paragraph()
p.style = 'Normal'
run = p.add_run(output_template)
run.font.name = 'Courier New'
run.font.size = Pt(9)

add_heading(doc, 'Output Format Notes', 2)
for item in [
    'The Documents Uploaded checklist always appears first.',
    'The Discrepancy Grid replaces the old full-field comparison table. It only includes Red and Yellow Flag items.',
    'Verified / OK: include only important fields (address, county, parcel, contract price) — not every possible field.',
    'Suggested Report Comment: only include when a flag clearly requires a report-level explanation (acreage conflict, partial conveyance, concessions, personal property, repairs, non-arms-length).',
    'Minor features (sheds, appliances, outbuildings) mentioned in MLS with no conflict → Information Only, not a flag.',
    'Deed/title: do not list as missing on every file. Only flag as missing when there is a legal, parcel, or conveyance conflict that cannot be resolved from uploaded documents.',
]:
    add_bullet(doc, item)

# ─── Section 9 ─────────────────────────────────────────────────────────────
add_heading(doc, '9. Standard Staff Prompt', 1)
add_para(doc, 'Anna can copy and paste this prompt after uploading the file packet:')

prompt_text = """Run the front-end consistency check on this appraisal write-up packet.

This check is being performed after the report was initially written up. Compare the report front page against the uploaded order form, contract, MLS subject sheet, tax/assessor record, parcel map, legal description, seller disclosure, and any other source documents.

Check property identity, address, city, county, ZIP, parcel number, legal description, site size/acreage, parcel count, buyer/borrower, seller/owner, client/lender, assignment type, FHA/VA case numbers, contract price, concessions, personal property, repairs/completion items, year built, HOA/PUD, utilities, contract execution, non-arms-length indicators, seller name vs. tax card owner, due date reasonableness, report fee reasonableness, investment property/rent schedule flags, and GP intended use/intended user.

Do not assume every mismatch is an error. Flag issues as Red Flag, Yellow Flag, Information Only, OK, or Unable to Verify. Return the required QC format with a documents uploaded checklist, discrepancy grid, recommended staff actions, and a suggested report comment only if useful and clearly supported."""

p = doc.add_paragraph()
run = p.add_run(prompt_text)
run.font.name = 'Courier New'
run.font.size = Pt(9)

# ─── Section 10 ─────────────────────────────────────────────────────────────
add_heading(doc, '10. Save and Rerun Rules', 1)
headers = ['Rule', 'Instruction']
rows = [
    ('Save every time', 'Save the GPT output in the workfile every time the check is run.'),
    ('File name', '[Property Address] - Write Up Review.pdf'),
    ('Red flags', 'Correct, document, or escalate before the file moves forward.'),
    ('Yellow flags', 'Review and resolve, or add a note explaining why no correction is needed.'),
    ('Rerun after corrections', 'Rerun the GPT after any corrections are made and save the updated review.'),
    ('Missing documents', 'If a needed document is missing, request it from the client or document that the item could not be verified.'),
]
add_table_with_header(doc, headers, rows)

# ─── Section 11 ─────────────────────────────────────────────────────────────
add_heading(doc, '11. Behavior Rules for the GPT', 1)
for item in [
    'Be concise and direct.',
    'Do not invent missing data.',
    'If text is unclear or unreadable, say exactly what could not be read.',
    'If documents conflict, show the conflicting source values.',
    'If a field is not found, say "Not Found" or "Unable to Verify."',
    'Do not assume a conflict is an error. Flag it and recommend review.',
    'Do not clear a red flag yourself.',
    'Do not make final legal, title, acreage, boundary, or ownership conclusions.',
    'Do not provide an appraisal value opinion.',
    'Do not perform comparable selection, market analysis, or adjustment review.',
    'Prioritize issues that could cause a report correction, revision request, client issue, or wrong front-end setup.',
    'Keep the output practical for Anna to use.',
]:
    add_bullet(doc, item)

# ─── Section 12 — NEW ───────────────────────────────────────────────────────
add_heading(doc, '12. Version History', 1)
headers = ['Version', 'Date', 'Changes']
rows = [
    ('1.0', 'Original', 'Initial release — 7 check categories.'),
    ('1.1', 'May 2026', 'Added: contract execution check, non-arms-length detection, seller name vs. tax card, utility consistency (Sec. 6F), GP intended use/user, report fee schedule (updated VA fees: SFR $700, MFH $750, 2-4 unit $800), due date reasonableness, rent schedule/investment property flag. Updated output format: Documents Uploaded checklist added; full field table replaced with Discrepancy Grid (flags only). Added seller disclosure to upload packet.'),
]
add_table_with_header(doc, headers, rows)

# ─── Save ───────────────────────────────────────────────────────────────────
import os
os.makedirs('/home/tommy/nabholz-tools/sop', exist_ok=True)
out = '/home/tommy/nabholz-tools/sop/Nabholz_Appraisal_Write-Up_Consistency_Check_SOP_v1.1.docx'
doc.save(out)
print(f'Saved: {out}')
